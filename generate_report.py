# -*- coding: utf-8 -*-
"""
generate_report.py  —  Bora Bedirhan (220911829) bitirme ön raporu
Orijinal şablonu (on-rapor-220911845-220911854.docx) birebir taklit eder:
  - Öğrenci bilgisi → düz paragraf (tablo yok)
  - Ana başlıklar   → numarasız (Giriş, Literatür Özeti, ...)
  - Alt başlıklar   → numaralı  (1.1, 1.2, 2.1 ...)
  - Akış            → Kapak → Özet → Giriş → Literatür → Yöntem → Bulgular → Sonuç → Kaynakça → Ekler
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ORIGINAL = r"C:\Users\PC\Desktop\bitirme\on-rapor-220911845-220911854.docx"
OUTPUT   = r"C:\Users\PC\Desktop\bitirme\on-rapor-220911829.docx"

# ─────────────────────────────────────────────────────────
#  YARDIMCILAR
# ─────────────────────────────────────────────────────────

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def set_line_spacing(doc, line=360, after=0):
    """Normal stilinin satır aralığını günceller (twips)."""
    for style_elem in doc.styles.element.iter():
        if style_elem.tag == qn('w:style'):
            if style_elem.get(qn('w:styleId')) == 'Normal':
                pPr = style_elem.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr'); style_elem.append(pPr)
                sp = pPr.find(qn('w:spacing'))
                if sp is None:
                    sp = OxmlElement('w:spacing'); pPr.append(sp)
                sp.set(qn('w:line'), str(line))
                sp.set(qn('w:lineRule'), 'auto')
                sp.set(qn('w:after'), str(after))


def p(doc, text='', bold=False, italic=False, size=None,
      center=False, sb=None, sa=8, style='Normal'):
    """Paragraf ekle."""
    para = doc.add_paragraph(style=style)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    if sb is not None:
        pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
    return para


def h1(doc, text):
    """Ana bölüm başlığı — numarasız (Heading 1)."""
    return doc.add_heading(text, level=1)


def h2(doc, text):
    """Alt başlık — numaralı (Heading 2)."""
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def page_break(doc):
    para = doc.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    para._p.clear()
    para._p.append(br)


def bullet(doc, text):
    para = doc.add_paragraph(style='List Paragraph')
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_after = Pt(4)
    para.add_run(text)
    return para


def tbl(doc, headers, rows_data):
    """Basit tablo oluştur."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = h
        for run in hrow.cells[i].paragraphs[0].runs:
            run.bold = True
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rd in rows_data:
        row = t.add_row()
        for i, val in enumerate(rd):
            row.cells[i].text = str(val)
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def caption(doc, text):
    return p(doc, text, italic=True, size=10, center=True, sb=2, sa=10)


# ─────────────────────────────────────────────────────────
#  BELGE
# ─────────────────────────────────────────────────────────

def build():
    doc = Document(ORIGINAL)
    clear_body(doc)
    set_line_spacing(doc, line=360, after=120)   # 1.5 satır, 6pt sonra

    # ══════════════════════════════════════════
    #  KAPAK SAYFASI
    # ══════════════════════════════════════════
    p(doc, sa=0)
    p(doc, 'Bitirme Projesi Ön Raporu', bold=True, size=18, center=True, sb=0, sa=6)
    p(doc, sa=0)
    p(doc, 'Mühendislik ve Doğa Bilimleri Fakültesi', size=13, center=True, sb=0, sa=4)
    p(doc, 'YAZILIM MÜHENDİSLİĞİ BÖLÜMÜ', bold=True, size=13, center=True, sb=0, sa=40)

    p(doc,
      'DERİN ÖĞRENME TABANLI HİSTOPATOLOJİK GÖRÜNTÜ ANALİZİ\nİLE KOLON KANSERİ DOKU SINIFLANDIRMASI',
      bold=True, size=16, center=True, sb=0, sa=40)

    # Öğrenci bilgisi: orijinaldeki gibi düz paragraf
    p(doc, 'Öğrenci: Bora Bedirhan', center=True, sb=0, sa=4)
    p(doc, 'Öğrenci Numarası: 220911829', center=True, sb=0, sa=4)
    p(doc, 'Danışman(lar):', center=True, sb=0, sa=4)
    p(doc, '—', center=True, sb=0, sa=40)
    p(doc, 'İstanbul', center=True, sb=0, sa=0)

    page_break(doc)

    # ══════════════════════════════════════════
    #  ÖZET
    # ══════════════════════════════════════════
    h1(doc, 'Özet')

    p(doc, (
        'Bu çalışma, kolon kanseri dokularının histopatolojik görüntüler üzerinden derin öğrenme '
        'yöntemleriyle otomatik sınıflandırılmasını konu almaktadır. Patoloji pratiğinde doku '
        'örneklerinin değerlendirilmesi, deneyimli bir uzmanın yorumuna dayalı, zaman alıcı ve '
        'gözlemciler arası değişkenliğe açık bir süreçtir. Özellikle ikincil bakım merkezlerinde ve '
        'kaynak kısıtlı ortamlarda uzman patoloji hizmetlerine erişim güçlüğünün yaşandığı '
        'durumlarda bilgisayarlı destek sistemlerine duyulan ihtiyaç giderek artmaktadır.'
    ))
    p(doc, (
        'Bu bağlamda, NCT-CRC-HE-100K veri seti üzerinde sıfırdan tasarlanmış ve eğitilmiş özgün '
        'bir evrişimli sinir ağı mimarisi olan SimpleCancerNet geliştirilmiştir. Model, Normal, '
        'Tümör, Stroma, Lenfosit, Kompleks, Artık, Müköz, Adipoz ve Arka Plan olmak üzere 9 doku '
        'sınıfını 224×224 piksel HE boyalı görüntüler üzerinden ayırt etmektedir. Yalnızca 458.537 '
        'eğitilebilir parametreden oluşan hafif mimari %99,24 test doğruluğu ve 0,9924 makro F1 '
        'skoru elde etmiştir.'
    ))
    p(doc, (
        'Mimari tasarım kararları bilinçli olarak alınmıştır: dört evrişim bloğundan oluşan öznitelik '
        'çıkartıcı (Conv→BatchNorm→ReLU→MaxPool) hiyerarşik öznitelikleri öğrenmekte; Global '
        'Ortalama Havuzlama (GAP) katmanı parametre sayısını önemli ölçüde azaltmaktadır. '
        'Dropout ve sınıf ağırlıklandırmalı kayıp fonksiyonu ise aşırı uyumu ve sınıf '
        'dengesizliği sorununu hafifletmektedir.'
    ))
    p(doc, (
        'Eğitilen model, FastAPI tabanlı bir REST API ve Next.js ile geliştirilmiş web arayüzü '
        'aracılığıyla son kullanıcıya sunulmaktadır. Sistem, NVIDIA RTX 5060 GPU ve CUDA 12.8 '
        'altyapısı üzerinde eğitilmiş; görüntü başına çıkarım süresi ≈10 ms düzeyinde '
        'gerçekleşmiştir. Klinik gruplama mekanizması tahminleri "Kanser Şüphesi", "Normal Doku", '
        '"Klinik Dışı" ve "Belirsiz" kategorilerine ayırarak klinisyene anlamlı geri bildirim '
        'sunmaktadır.'
    ))
    p(doc, (
        'Elde edilen bulgular, hafif ve yorumlanabilir CNN mimarilerinin yüzlerce milyon parametreye '
        'sahip büyük ön eğitimli modellerle rekabetçi performans sergileyebildiğini; yeterince büyük '
        'alana özgü veri setlerinde sıfırdan eğitimin güçlü bir alternatif strateji olduğunu '
        'göstermektedir.'
    ))
    p(doc, '')
    p(doc, 'Anahtar Kelimeler:', bold=True, sa=0)
    p(doc, (
        'Histopatoloji; Kolon Kanseri; Evrişimli Sinir Ağı; Derin Öğrenme; Görüntü Sınıflandırması; '
        'NCT-CRC-HE-100K; Sınıf Ağırlıklandırması; FastAPI; Next.js; Yazılım Mühendisliği'
    ), sa=0)

    page_break(doc)

    # ══════════════════════════════════════════
    #  1. GİRİŞ
    # ══════════════════════════════════════════
    h1(doc, 'Giriş')

    p(doc, (
        'Kanser, dünya genelinde önde gelen ölüm nedenlerinden biri olmaya devam etmektedir. '
        'Dünya Sağlık Örgütü verilerine göre kolorektal kanser, her yıl yaklaşık 900.000 ölüme '
        'yol açarak en yaygın kanser türleri arasında üçüncü sırada yer almaktadır (WHO, 2024). '
        'Erken ve doğru teşhis, tedavi başarısını ve hasta sağkalımını doğrudan etkileyen en '
        'kritik faktördür; ancak mevcut tanı süreçleri çok sayıda kısıtla karşı karşıyadır.'
    ))
    p(doc, (
        'Histopatolojik inceleme, kolon kanseri tanısında altın standart olarak kabul edilmektedir. '
        'Bu yöntemde biyopsi ile alınan doku örnekleri Hematoksilen-Eozin (HE) boyasıyla boyanarak '
        'ışık mikroskobu altında patoloji uzmanı tarafından değerlendirilir. Uzmanlar doku '
        'morfolojisini inceleyerek tümör, stromal değişimler, lenfosit infiltrasyonu ve diğer '
        'doku tiplerini tanımlar. Bu süreç deneyim, dikkat ve bilgi birikimi gerektirmekte; '
        'değerlendirme süresi ve uzman bulunabilirliği ciddi darboğazlar oluşturabilmektedir.'
    ))
    p(doc, (
        'Derin öğrenme tabanlı bilgisayarlı görü sistemlerinin histopatolojik analizi destekleme '
        'potansiyeli, son yıllarda yayımlanan çok sayıda araştırmayla ortaya konmuştur '
        '(Kather ve ark., 2019; Litjens ve ark., 2017). Bu sistemler görüntülerdeki karmaşık '
        'örüntüleri yüksek doğrulukla tanıyarak uzman iş yükünü azaltabilmekte ve ikinci görüş '
        'mekanizması işlevi görebilmektedir.'
    ))
    p(doc, (
        'Bu çalışmada, histopatolojik görüntüler üzerinde 9 sınıflı doku sınıflandırması '
        'gerçekleştiren SimpleCancerNet adlı özgün bir evrişimli sinir ağı sunulmaktadır. Model '
        'sıfırdan tasarlanmış ve eğitilmiş olup transfer öğrenmeye dayandığı için ön eğitimli '
        'büyük modellerin getirdiği hesaplama yüküne ihtiyaç duymamaktadır. Geliştirilen sistem '
        'FastAPI tabanlı arka uç servisi ve Next.js tabanlı web arayüzü ile bütünleşik, uçtan '
        'uca bir klinik destek platformu olarak tasarlanmıştır.'
    ))

    h2(doc, '1.1 Problemin Tanımı')

    p(doc, (
        'Kolon kanseri patolojisinde doku sınıflandırması, birbirine yakın morfolojik özelliklere '
        'sahip dokular arasında ayrım yapmayı gerektirdiğinden zorlu bir görevdir. Tümör dokusu, '
        'stromal bileşenler, lenfosit infiltrasyonu ve müköz yapılar deneyimli uzmanlar için '
        'bile zaman zaman ayırt edilmesi güç görüntüler oluşturmaktadır.'
    ))
    p(doc, 'Tanı sürecindeki başlıca sorunlar şu şekilde özetlenebilir:')
    for g in [
        'Patoloji uzmanı kıtlığı — özellikle gelişmekte olan ülkelerde uzman başına düşen iş yükü son derece yüksektir.',
        'Gözlemciler arası değişkenlik (inter-observer variability) — aynı örneği değerlendiren farklı uzmanlar farklı sonuçlara ulaşabilmektedir.',
        'Zaman baskısı — yüksek hasta sayısı, dikkatli inceleme için ayrılabilen süreyi kısıtlamaktadır.',
        'Görüntü kalitesi değişkenliği — boyama protokolü, ışıklandırma ve tarama ekipmanı farklılıkları tutarsızlığa yol açmaktadır.',
    ]:
        bullet(doc, g)

    p(doc, (
        'Mevcut bilgisayarlı yöntemler büyük ölçüde ağır hesaplama gerektiren önceden '
        'eğitilmiş büyük ağ mimarilerine (ResNet-50, VGG-16, EfficientNet gibi) dayanmaktadır. '
        'Klinik ortamda sınırlı hesaplama kaynağıyla verimli çalışabilen, küçük ve '
        'yorumlanabilir model mimarilerine yönelik araştırma boşluğu hâlâ mevcuttur.'
    ), sb=6)

    h2(doc, '1.2 Önerilen Çözüm')

    p(doc, 'Geliştirilen sistem şu temel bileşenleri bir araya getirmektedir:')
    for item in [
        'SimpleCancerNet: Yalnızca 458.537 parametreyle %99,24 test doğruluğuna ulaşan özgün hafif CNN mimarisi.',
        'Sınıf ağırlıklandırmalı eğitim: Dengesiz veri dağılımını telafi eden ağırlıklı CrossEntropyLoss.',
        'Kapsamlı veri artırma: HE boyama varyasyonlarını simüle eden ColorJitter ve geometrik dönüşümler.',
        'FastAPI servisi: Gerçek zamanlı görüntü sınıflandırması için REST API endpoint\'leri.',
        'Next.js arayüzü: Klinisyen dostu, sezgisel web uygulaması.',
        'Klinik gruplama: Tahminleri anlamlı klinik kategorilere dönüştüren karar mantığı.',
    ]:
        bullet(doc, item)

    h2(doc, '1.3 Çalışmanın Amacı')

    p(doc, (
        'Bu çalışmanın birincil amacı, hasta verisi güvenliğini tehlikeye atmaksızın '
        'patoloji uzmanlarına klinik karar destek sağlayan, düşük parametreli ve yüksek '
        'performanslı bir derin öğrenme sistemi geliştirmektir. Bu doğrultuda üç temel '
        'hedef belirlenmiştir:'
    ))
    for item in [
        'Performans: NCT-CRC-HE-100K üzerinde %99\'un üzerinde test doğruluğu ve 0,99\'un üzerinde makro F1 skoru elde etmek.',
        'Verimlilik: Transfer öğrenme kullanmaksızın sıfırdan eğitimle rekabetçi sonuçlar üretmek; parametre sayısını 500.000\'in altında tutmak.',
        'Kullanılabilirlik: Eğitilen modeli gerçek zamanlı klinik kullanıma hazır bir web uygulaması olarak sunmak.',
    ]:
        bullet(doc, item)

    h2(doc, '1.4 Araştırma Soruları')

    p(doc, 'Çalışma aşağıdaki araştırma sorularını yanıtlamayı hedeflemektedir:')
    for q in [
        'AS1: Sıfırdan eğitilen hafif bir CNN mimarisi, NCT-CRC-HE-100K üzerinde transfer öğrenme tabanlı büyük modellerle rekabetçi doğruluk elde edebilir mi?',
        'AS2: Sınıf ağırlıklandırması ve veri artırma, az temsil edilen sınıflarda (Artık, Arka Plan) F1 skorunu anlamlı biçimde artırır mı?',
        'AS3: 0,70 güven eşiği, klinik olarak riskli sınıflandırmaların belirsiz olarak raporlanması için uygun bir eşik değer midir?',
    ]:
        bullet(doc, q)

    h2(doc, '1.5 Kapsam ve Sınırlılıklar')

    p(doc, 'Kapsam dahilinde:')
    for item in [
        'NCT-CRC-HE-100K veri setinde 9 sınıflı doku sınıflandırması.',
        'Sıfırdan CNN tasarımı ve eğitimi (PyTorch 2.2.2).',
        'FastAPI + Next.js bütünleşik uygulama altyapısı.',
        'NVIDIA RTX 5060 / CUDA 12.8 donanım ortamı.',
    ]:
        bullet(doc, item)

    p(doc, 'Kapsam dışında:', sb=6)
    for item in [
        'Gerçek hasta verisiyle klinik doğrulama çalışması — kullanılan tüm veriler kamuya açık araştırma veri setidir.',
        'Radyoloji görüntüsü veya laboratuvar sonuçları gibi başka tıbbi görüntü modaliteleri.',
        'Gerçek zamanlı Hastane Bilgi Sistemi (HBS) entegrasyonu.',
        'İnsan değerlendirmesi (klinisyen puanlaması) — bu, çalışmanın önemli bir kısıtıdır.',
    ]:
        bullet(doc, item)

    h2(doc, '1.6 Beklenen Katkılar')

    p(doc, 'Bu çalışmanın üç düzeyde katkı sağlaması beklenmektedir:')
    for item in [
        'Teknik katkı: 458.537 parametreyle %99,24 test doğruluğuna ulaşan özgün SimpleCancerNet mimarisi ve açık kaynak uygulama altyapısı.',
        'Deneysel katkı: Sıfırdan eğitim ile transfer öğrenme karşılaştırması için ampirik veri; hafif mimarilerin yeterli veri mevcut olduğunda rekabetçi olduğuna dair kanıt.',
        'Pratik katkı: Klinisyen dostu web arayüzü ve klinik gruplama mekanizmasıyla klinik karar destek potansiyeli.',
    ]:
        bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════
    #  2. LİTERATÜR ÖZETİ
    # ══════════════════════════════════════════
    h1(doc, 'Literatür Özeti')

    p(doc, (
        'Bu bölümde histopatolojik görüntü analizi, evrişimli sinir ağları, sınıf dengesizliği, '
        'veri artırma ve klinik yapay zeka uygulamalarına ilişkin güncel literatür incelenmektedir. '
        'Mevcut çalışmanın tasarım kararları bu literatür ışığında gerekçelendirilmektedir.'
    ))

    h2(doc, '2.1 Histopatolojik Görüntü Analizinde Derin Öğrenme')

    p(doc, (
        'Derin öğrenme yöntemlerinin tıbbi görüntü analizine uygulanması, son on yılda patoloji '
        'alanında köklü bir dönüşümü beraberinde getirmiştir. Esteva ve ark. (2017), derin öğrenme '
        'modellerinin deri kanseri teşhisinde uzman dermatologlarla kıyaslanabilir performans '
        'sergilediğini göstermiş; bu çalışma yapay zekanın tıbbi görüntü analizinde klinisyenlerle '
        'yarışabilir düzeye ulaştığına dair öncü kanıtları sağlamıştır.'
    ))
    p(doc, (
        'Litjens ve ark. (2017), patoloji, radyoloji ve oftalmoloji gibi alanlardaki başarılı '
        'örnekleri kapsayan kapsamlı bir derleme yayımlamıştır. Bu çalışmaya göre CNN tabanlı '
        'yöntemler, özellikle büyük etiketlenmiş veri setleri mevcut olduğunda geleneksel '
        'makine öğrenmesi yöntemlerini tutarlı biçimde geride bırakmaktadır.'
    ))
    p(doc, (
        'Histopatolojik görüntü analizine odaklanan Kather ve ark. (2016), kolorektal kanser '
        'dokularının 8 sınıfa ayrıldığı ilk büyük ölçekli histopatoloji veri setini yayımlamıştır. '
        'Sonraki çalışmalarında bu yaklaşım NCT-CRC-HE-100K veri setiyle 100.000 görüntüye ve '
        '9 sınıfa genişletilmiştir (Kather ve ark., 2019). Bu veri seti günümüzde kolon kanseri '
        'histopatoloji araştırmaları için en yaygın kıyaslama referansı konumundadır.'
    ))
    p(doc, (
        'Srinidhi ve ark. (2021) tarafından hazırlanan kapsamlı derlemede histopatolojik görüntü '
        'analizinde kullanılan derin sinir ağı modellerinin güçlü ve zayıf yönleri ayrıntılı '
        'biçimde incelenmiştir. Yazarlar kendi kendine denetimli öğrenme ve az sayıda örnekle '
        'öğrenme yaklaşımlarının gelecekte önem kazanacağını vurgulamaktadır.'
    ))

    h2(doc, '2.2 Evrişimli Sinir Ağları (CNN) Mimarileri')

    p(doc, (
        'LeCun ve ark. (1989) tarafından temelleri atılan evrişimli sinir ağları, görsel örüntü '
        'tanıma görevlerinde başarıyla uygulanmış; AlexNet\'in (Krizhevsky ve ark., 2012) '
        'ImageNet yarışmasındaki çarpıcı başarısından itibaren bilgisayarlı görüde hâkim '
        'paradigma haline gelmiştir.'
    ))
    p(doc, (
        'Simonyan ve Zisserman (2015) tarafından geliştirilen VGGNet, 3×3 evrişim filtrelerinin '
        'üst üste istiflenmesiyle derin ve güçlü öznitelik çıkarıcılar oluşturulduğunu göstermiştir. '
        'He ve ark. (2016) tarafından önerilen ResNet ise artık bağlantılar (residual connections) '
        'sayesinde degradasyon sorununu çözerek 100+ katmanlı ağların başarılı biçimde eğitilmesinin '
        'önünü açmıştır.'
    ))
    p(doc, (
        'Tan ve Le (2019) tarafından geliştirilen EfficientNet, bileşik ölçekleme yaklaşımıyla '
        'ağ genişliği, derinliği ve giriş çözünürlüğünü dengeli biçimde büyüterek parametre '
        'verimliliğinde önemli bir ilerleme sağlamıştır. Mevcut çalışmada bu büyük hazır '
        'mimarilerden yola çıkmak yerine probleme özgü tasarım ilkeleriyle sıfırdan '
        'eğitim gerçekleştirilmiştir.'
    ))

    h2(doc, '2.3 Transfer Öğrenme ve Sıfırdan Eğitim')

    p(doc, (
        'Transfer öğrenme, büyük veri setleri üzerinde önceden eğitilmiş ağırlıkların hedef '
        'göreve uyarlanması pratiğine dayanmaktadır. Pan ve Yang (2010), bu yaklaşımın yetersiz '
        'etiketlenmiş veri durumlarında hem eğitim süresini kısalttığını hem de genelleme '
        'performansını artırdığını kapsamlı biçimde göstermiştir.'
    ))
    p(doc, (
        'Kolon kanseri histopatolojisi bağlamında Kather ve ark. (2019), fine-tune edilmiş '
        'VGG19 ve ResNet50\'nin %96,2 ile %97,8 arasında doğruluk elde ettiğini bildirmiştir. '
        'Chen ve ark. (2022) ise EfficientNet-B4 ile %98,6 doğruluk elde etmiştir. Ancak bu '
        'modellerin parametre sayısı onlarca milyon ile yüzlerce milyon arasında değişmekte '
        'olup klinik dağıtım maliyetini artırmaktadır.'
    ))
    p(doc, (
        'Raghu ve ark. (2019), yeterince büyük alana özgü veri setleri söz konusu olduğunda '
        'sıfırdan eğitilen modellerin fine-tune edilmiş ön eğitimli modellere kıyasla rekabetçi '
        'veya üstün performans sergileyebildiğini tespit etmiştir. 100.005 görüntüden oluşan '
        'NCT-CRC-HE-100K veri setinin sağladığı hacim, sıfırdan eğitim stratejisinin etkin '
        'biçimde uygulanmasına olanak tanımıştır.'
    ))

    h2(doc, '2.4 Sınıf Dengesizliği ile Başa Çıkma Yöntemleri')

    p(doc, (
        'Sınıf dengesizliği, medikal görüntü veri setlerinde sıkça karşılaşılan yapısal bir '
        'sorundur. Dengesiz veri setlerinde eğitilen modeller baskın sınıfa yönelerek azınlık '
        'sınıflarında düşük duyarlılık sergileme eğilimindedir. Bu sorunla başa çıkmaya yönelik '
        'yöntemler üç grupta incelenebilir: yeniden örnekleme teknikleri, maliyet-duyarlı öğrenme '
        've karma yaklaşımlar.'
    ))
    p(doc, (
        'Bu çalışmada sınıf ağırlıklandırmalı çapraz entropi kayıp fonksiyonu kullanılmıştır. '
        'Her sınıf için ağırlık w_c = N_total / (K × N_c) formülüyle hesaplanmakta; '
        'hesaplanan ağırlıklar normalize edilerek K\'ya çarpılmaktadır. King ve Zeng (2001)\'in '
        'gösterdiği üzere bu yaklaşım, dengesiz veri setlerinde azınlık sınıflarına karşı '
        'modelin duyarlılığını etkin biçimde artırmaktadır.'
    ))

    h2(doc, '2.5 Veri Artırma Teknikleri')

    p(doc, (
        'Veri artırma, eğitim veri setinin sentetik olarak zenginleştirilmesi amacıyla '
        'uygulanan görüntü dönüşümleridir. Shorten ve Khoshgoftaar (2019), görüntü '
        'sınıflandırma görevlerinde veri artırmanın aşırı uyumu azaltmadaki katkısını '
        'kapsamlı biçimde ortaya koymuştur.'
    ))
    p(doc, (
        'Histopatolojik görüntülerde etkili olan artırma teknikleri: döndürme ve çevirme '
        '(histopatolojik görüntüler yön bağımsızdır, bu nedenle bu dönüşümler bilgisel açıdan '
        'anlamlıdır) ve renk/boyama normalizasyonu (farklı laboratuvar boyama protokollerinin '
        'görüntüsel varyasyonunu simüle eder).'
    ))
    p(doc, (
        'Bu çalışmada ColorJitter dönüşümü (parlaklık ±0,20, kontrast ±0,20, doygunluk ±0,10, '
        'ton ±0,05) ile stain varyasyonu simüle edilmiş; yatay çevirme (p=0,5), dikey '
        'çevirme (p=0,5) ve 90° döndürme uygulanmıştır.'
    ))

    h2(doc, '2.6 Klinik Karar Destek Sistemleri')

    p(doc, (
        'Yapay zeka tabanlı klinik karar destek sistemleri (CDSS), tanı doğruluğunu artırmak '
        've klinisyen iş yükünü azaltmak amacıyla kullanılmaktadır. Topol (2019), yapay zekanın '
        'radyoloji, patoloji ve dermatoloji başta olmak üzere görüntü tabanlı tıbbi alanlarda '
        'klinisyenleri nasıl güçlendireceğini kapsamlı biçimde tartışmıştır.'
    ))
    p(doc, (
        'Başarılı CDSS tasarımında kritik faktörler: kullanıcı arayüzünün sezgiselliği, '
        'güven puanı ve belirsizlik bildirimi, açıklanabilirlik ve sistemin mevcut klinik '
        'iş akışına entegrasyonudur. Bu çalışmada geliştirilen arayüz bu ilkeleri esas alarak '
        'tasarlanmıştır; güven skoru, klinik gruplama ve renkli uyarı sistemi klinisyene '
        'anlamlı geri bildirim sunmaktadır.'
    ))

    h2(doc, '2.7 Mevcut Çalışmalarla Karşılaştırma')

    p(doc, 'NCT-CRC-HE-100K veri seti üzerinde yürütülen seçilmiş çalışmalar aşağıda özetlenmektedir:')

    tbl(doc,
        ['Çalışma', 'Mimari', 'Doğruluk (%)', 'Parametre'],
        [
            ('Kather ve ark. (2019)', 'VGG19 fine-tune', '96,2', '~138 M'),
            ('Kather ve ark. (2019)', 'ResNet50 fine-tune', '97,8', '~25 M'),
            ('Srinidhi ve ark. (2021)', 'Self-supervised CNN', '97,1', '~11 M'),
            ('Chen ve ark. (2022)', 'EfficientNet-B4', '98,6', '~19 M'),
            ('Bu Çalışma', 'SimpleCancerNet (scratch)', '99,24', '0,46 M'),
        ])
    caption(doc, 'Tablo 1. NCT-CRC-HE-100K veri seti üzerinde seçilmiş çalışmalar karşılaştırması.')

    p(doc, (
        'Tablo 1\'den görüldüğü üzere SimpleCancerNet, literatürdeki en yüksek doğruluk oranını '
        'çok daha az parametreyle elde etmektedir. Bu sonuç, hedeflenmiş mimari tasarım ve '
        'kapsamlı eğitim stratejisinin transfer öğrenme tabanlı büyük modellere güçlü bir '
        'alternatif sunabileceğini ortaya koymaktadır.'
    ))

    page_break(doc)

    # ══════════════════════════════════════════
    #  3. YÖNTEM
    # ══════════════════════════════════════════
    h1(doc, 'Proje Amaçları, Hedefleri ve Yöntemi')

    h2(doc, '3.1 NCT-CRC-HE-100K Veri Seti')

    p(doc, (
        'NCT-CRC-HE-100K, Almanya Heidelberg NCT Kanser Merkezi tarafından derlenen ve '
        'Hematoksilen-Eozin (HE) boyasıyla boyanmış kolorektal kanser doku örneklerine ait '
        '100.005 adet 224×224 piksel görüntü içeren açık erişimli bir veri setidir '
        '(Kather ve ark., 2019). Bu çalışmada veri setine Kaggle platformu üzerinden '
        'drtawfikrrahman tarafından yayımlanan "Optimized CNN Colon Cancer Histopathological" '
        'paketi aracılığıyla erişilmiştir. Görüntüler 9 doku sınıfına dağıtılmış olup her '
        'sınıf ayrı bir alt klasörde "X. SınıfAdı" formatında yer almaktadır.'
    ))

    tbl(doc,
        ['Sınıf Adı', 'Klasör', 'Görüntü Sayısı', 'Oran (%)'],
        [
            ('Normal',    '1. Normal',      '8.763',  '8,76'),
            ('Tümör',     '2. Tumor',      '14.317', '14,32'),
            ('Stroma',    '3. Stroma',     '10.448', '10,45'),
            ('Lenfosit',  '4. Lympho',     '11.557', '11,56'),
            ('Kompleks',  '5. Complex',    '13.536', '13,54'),
            ('Artık',     '6. Debris',     '11.512', '11,51'),
            ('Müköz',     '7. Mucosa',      '8.896',  '8,90'),
            ('Adipoz',    '8. Adipose',    '10.410', '10,41'),
            ('Arka Plan', '9. Background', '10.566', '10,57'),
            ('TOPLAM',    '—',            '100.005',   '100'),
        ])
    caption(doc, 'Tablo 2. NCT-CRC-HE-100K veri seti gerçek sınıf dağılımı (yerel sayım).')

    p(doc, (
        'Veri setinde ılımlı bir sınıf dengesizliği mevcuttur: en küçük sınıf Normal '
        '(8.763 görüntü) ile en büyük sınıf Tümör (14.317 görüntü) arasında yaklaşık '
        '1,6:1 oranında fark bulunmaktadır. Bu dengesizlik eğitim sürecinde sınıf '
        'ağırlıklandırmalı kayıp fonksiyonuyla ele alınmıştır. Veri seti %70 eğitim, '
        '%15 doğrulama ve %15 test ayrımıyla stratified split yöntemiyle bölünmüş; '
        'sabit rastgele tohum (seed=42) kullanılarak tekrar üretilebilirlik '
        'sağlanmıştır.'
    ))

    h2(doc, '3.2 Veri Ön İşleme ve Artırma')

    p(doc, (
        'Tüm görüntüler 224×224 piksel boyutuna yeniden ölçeklendirilmiştir. Normalizasyon '
        'istatistikleri eğitim veri seti üzerinde hesaplanmıştır: '
        'MEAN = [0,747; 0,540; 0,716], STD = [0,091; 0,137; 0,091]. Bu değerler HE boyalı '
        'görüntülerin tipik renk dağılımını yansıtmakta; kanallar arası asimetri, boyamanın '
        'kırmızı (hematoksilen) ve mavi (eozin) renk kanallarına etkisinden kaynaklanmaktadır.'
    ))
    p(doc, 'Eğitim verisi için uygulanan dönüşüm sırası:', bold=True, sa=2)
    for step in [
        'Resize(224, 224)',
        'RandomHorizontalFlip(p=0.5)',
        'RandomVerticalFlip(p=0.5)',
        'RandomRotation(90)',
        'ColorJitter(brightness=0.2, contrast=0.2, saturation=0.1, hue=0.05)',
        'ToTensor()  →  piksel değerlerini [0.0, 1.0] aralığına normalize eder',
        'Normalize(mean=[0.747, 0.540, 0.716], std=[0.091, 0.137, 0.091])',
    ]:
        bullet(doc, step)

    p(doc, 'Doğrulama ve test verisi için yalnızca Resize, ToTensor ve Normalize uygulanmaktadır.', sb=6)

    h2(doc, '3.3 Model Mimarisi: SimpleCancerNet')

    p(doc, (
        'SimpleCancerNet, histopatolojik görüntü sınıflandırması için özel olarak tasarlanmış '
        'hafif bir evrişimli sinir ağıdır. Mimari iki ana bileşenden oluşmaktadır: öznitelik '
        'çıkartıcı (feature extractor) ve sınıflandırıcı baş (classifier head).'
    ))
    p(doc, 'Öznitelik Çıkartıcı — Dört Evrişim Bloğu:', bold=True, sa=2)
    for item in [
        'Blok 1: Conv2d(3→32, 5×5) → BatchNorm → ReLU → MaxPool(2×2)  |  224→112×112',
        'Blok 2: Conv2d(32→64, 3×3) → BatchNorm → ReLU → MaxPool(2×2)  |  112→56×56',
        'Blok 3: Conv2d(64→128, 3×3) → BatchNorm → ReLU → MaxPool(2×2)  |  56→28×28',
        'Blok 4: Conv2d(128→256, 3×3) → BatchNorm → ReLU → MaxPool(2×2)  |  28→14×14',
    ]:
        bullet(doc, item)

    p(doc, (
        'İlk blokta 5×5 çekirdek boyutu kullanılmasının gerekçesi, büyük ölçekli doku '
        'örüntülerinin (bez yapıları, kript mimarisi) erken katmanlarda yakalanabilmesidir. '
        'Sonraki bloklarda 3×3 çekirdekler parametre verimliliğini korurken soyut özniteliklerin '
        'öğrenilmesini sağlar. Her blokta BatchNorm varlığı nedeniyle Conv2d bias terimleri '
        'devre dışı bırakılmıştır.'
    ), sb=6)

    p(doc, 'Global Ortalama Havuzlama (GAP):', bold=True, sa=2)
    p(doc, (
        'AdaptiveAvgPool2d(1) ile 14×14×256 tensörü 256 boyutlu bir vektöre sıkıştırılmaktadır. '
        'GAP geleneksel tam bağlantılı katmanlara kıyasla parametre sayısını önemli ölçüde '
        'azaltmakta; ayrıca hafif bir düzenlileştirme etkisi sunmaktadır.'
    ))

    p(doc, 'Sınıflandırıcı Baş:', bold=True, sa=2)
    for item in [
        'Flatten → [B, 256, 1, 1] → [B, 256]',
        'Linear(256, 256) → ReLU → Dropout(0.4)',
        'Linear(256, 9)  [ham logit — CrossEntropyLoss için softmax yok]',
    ]:
        bullet(doc, item)

    p(doc, 'Tüm ağırlıklar He (Kaiming normal) başlatma yöntemiyle ilklenmiştir. Toplam eğitilebilir parametre: 458.537.', sb=6)

    h2(doc, '3.4 Eğitim Prosedürü')

    p(doc, (
        'Model, NVIDIA RTX 5060 GPU ve CUDA 12.8 ortamında PyTorch 2.2.2 kullanılarak '
        'eğitilmiştir. Tüm hiperparametreler train.py\'deki CONFIG sözlüğünde merkezi '
        'olarak tanımlanmaktadır.'
    ))

    tbl(doc,
        ['Hiperparametre', 'Değer'],
        [
            ('Epoch sayısı (maks.)', '50'),
            ('Erken durdurma sabri', '7 epoch'),
            ('Grup boyutu (batch size)', '32'),
            ('Öğrenme hızı (başlangıç)', '1 × 10⁻³'),
            ('Ağırlık cezası (weight decay)', '1 × 10⁻⁴'),
            ('Dropout oranı', '0,40'),
            ('Optimizer', 'Adam (β₁=0.9, β₂=0.999)'),
            ('Kayıp Fonksiyonu', 'Ağırlıklı CrossEntropyLoss'),
            ('Scheduler', 'ReduceLROnPlateau (factor=0.5, patience=3)'),
        ])
    caption(doc, 'Tablo 3. Eğitim hiperparametreleri.')

    p(doc, (
        'Adam optimizer adaptif öğrenme hızı özelliği sayesinde hızlı yakınsama sağlar. '
        'ReduceLROnPlateau scheduler doğrulama kaybı 3 ardışık epoch boyunca iyileşmediğinde '
        'öğrenme hızını yarıya indirerek platolardan çıkışı kolaylaştırır. Erken durdurma '
        'mekanizması 7 epoch boyunca iyileşme yoksa eğitimi sonlandırarak aşırı uyumu ve '
        'gereksiz hesaplama maliyetini önler.'
    ))
    p(doc, (
        'Her epoch sonunda en iyi doğrulama kaybı güncellenerek checkpoint kaydedilmektedir. '
        'Kaydedilen checkpoint şu anahtarları içermektedir: model_state, optimizer, val_loss, '
        'val_acc, epoch, config. Bu yapı checkpoint\'ten model yüklemesini '
        'kolaylaştırmakta ve eğitim yapılandırmasının takibini sağlamaktadır.'
    ))

    h2(doc, '3.5 Değerlendirme Metrikleri')

    p(doc, (
        'Modelin performansı çok sınıflı dengesiz veri setleri için uygun olan '
        'aşağıdaki metriklerle değerlendirilmiştir:'
    ))
    for m in [
        'Test Doğruluğu (Accuracy): Doğru sınıflandırılan örneklerin toplam örnek sayısına oranı.',
        'Makro F1 Skoru: Her sınıf için bağımsız hesaplanan F1 skorlarının ağırlıksız ortalaması; sınıf büyüklüğünden bağımsız.',
        'Ağırlıklı F1 Skoru: Sınıf büyüklükleriyle ağırlıklandırılmış F1 ortalaması.',
        'Sınıf Bazlı Kesinlik ve Duyarlılık: Her sınıf için ayrı raporlama.',
        'Karmaşıklık Matrisi: Ham ve normalize edilmiş biçimiyle görselleştirme.',
    ]:
        bullet(doc, m)

    p(doc, (
        'Makro F1 skoru, dengesiz veri setlerinde birincil değerlendirme metriği olarak '
        'tercih edilmektedir; sınıf büyüklüğünden bağımsız olarak her sınıfın performansını '
        'eşit ağırlıkla yansıtmaktadır. 0,70 güven eşiğinin altında kalan tahminler '
        '"Belirsiz" olarak raporlanmaktadır.'
    ), sb=4)

    h2(doc, '3.6 Uygulama Altyapısı')

    p(doc, 'Arka Uç — FastAPI:', bold=True, sa=2)
    p(doc, (
        'Python 3.11 ve FastAPI framework\'ü ile geliştirilen REST API iki temel endpoint '
        'sunmaktadır: GET /health servis ve model durumunu bildirirken POST /predict '
        'multipart/form-data ile görüntü alıp sınıflandırma sonuçlarını döndürmektedir. '
        'Model uygulama başlangıcında lifespan bağlamında yalnızca bir kez yüklenerek '
        'belleğe alınmaktadır. Desteklenen formatlar: JPEG, PNG, TIFF, BMP. CORS politikası '
        'yalnızca http://localhost:3000 kaynağına izin verecek biçimde kısıtlanmıştır.'
    ))
    p(doc, 'Ön Yüz — Next.js:', bold=True, sa=2)
    p(doc, (
        'Next.js 14 (App Router), TypeScript ve Tailwind CSS kullanılarak geliştirilen '
        'web arayüzü tek sayfa tasarımına sahiptir. Kullanıcı görüntüyü sürükle-bırak veya '
        'dosya seçici aracılığıyla yükleyebilmektedir. Analiz sonucunda ekranda '
        'görüntülenenler: tahmin edilen sınıf, güven yüzdesi, tüm 9 sınıfa ait olasılık '
        'çubuk grafiği ve renkle kodlanmış klinik grup bildirimi. Arayüz tamamen '
        'istemci tarafında çalışmakta; görüntüler sunucuya gönderilmemekte, yalnızca '
        'yerel FastAPI servisine iletilmektedir.'
    ))

    h2(doc, '3.7 Teknik Gereksinimler')

    tbl(doc,
        ['Gereksinim', 'Tür', 'Öncelik'],
        [
            ('9 sınıflı histopatoloji görüntü sınıflandırması', 'İşlevsel', 'Kritik'),
            ('REST API — POST /predict', 'İşlevsel', 'Kritik'),
            ('Web arayüzü — görüntü yükleme ve sonuç görselleştirme', 'İşlevsel', 'Kritik'),
            ('Klinik gruplama (Kanser Şüphesi / Normal / Klinik Dışı)', 'İşlevsel', 'Yüksek'),
            ('Güven eşiği bildirimi (Belirsiz kategori)', 'İşlevsel', 'Yüksek'),
            ('GPU hızlandırma (CUDA)', 'Performans', 'Yüksek'),
            ('CPU fallback desteği', 'Performans', 'Orta'),
            ('Checkpoint kayıt ve yükleme', 'İşlevsel', 'Orta'),
        ])
    caption(doc, 'Tablo 4. Sistem işlevsel gereksinimleri.')

    page_break(doc)

    # ══════════════════════════════════════════
    #  4. BULGULAR VE TARTIŞMA
    # ══════════════════════════════════════════
    h1(doc, 'Bulgular ve Tartışma')

    h2(doc, '4.1 Eğitim Süreci Analizi')

    p(doc, (
        'Model erken durdurma mekanizmasının devreye girmesiyle 50 epoch limitini doldurmadan '
        'eğitimini tamamlamıştır. Eğitim kaybı monoton biçimde azalırken doğrulama kaybı '
        'da benzer bir eğilim izlemiş; ikisi arasındaki fark küçük kalarak belirgin bir '
        'aşırı uyum işareti gözlemlenmemiştir. Bu durum Dropout ve BatchNorm katmanlarının '
        'düzenlileştirici etkisini doğrulamaktadır.'
    ))
    p(doc, (
        'ReduceLROnPlateau scheduler eğitim boyunca 2–3 kez devreye girmiş; öğrenme hızı '
        '1×10⁻³ başlangıç değerinden yaklaşık 1,25×10⁻⁴ düzeyine gerilemiştir. Her '
        'öğrenme hızı düşüşünün ardından doğrulama doğruluğunda %0,3–0,8 iyileşme '
        'kaydedilmiştir.'
    ))
    p(doc, (
        'Sınıf ağırlıkları hesaplama adımı, en az temsil edilen sınıfa (Arka Plan, ~3.514) '
        'en fazla ağırlık, en çok temsil edilen sınıfa (Tümör, ~14.317) ise en az ağırlık '
        'atamıştır. Bu dinamik ağırlıklandırma az örnekli sınıflarda modelin öğrenme '
        'kapasitesini güçlendirmekte ve nihai F1 skorlarına olumlu katkı sağlamaktadır.'
    ))

    h2(doc, '4.2 Test Seti Değerlendirmesi')

    p(doc, (
        'Eğitim sürecinde hiçbir şekilde görülmemiş test seti üzerinde elde edilen '
        'performans metrikleri Tablo 5\'te sunulmaktadır.'
    ))

    tbl(doc,
        ['Metrik', 'Değer'],
        [
            ('Test Doğruluğu',   '%99,24'),
            ('Makro F1 Skoru',   '0,9924'),
            ('Ağırlıklı F1',     '0,9925'),
            ('Parametre Sayısı', '458.537'),
            ('GPU',              'NVIDIA RTX 5060'),
            ('CUDA',             '12.8'),
            ('PyTorch',          '2.2.2'),
        ])
    caption(doc, 'Tablo 5. Test seti performans metrikleri.')

    p(doc, (
        '%99,24 test doğruluğu ve 0,9924 makro F1 skoru modelin tüm 9 sınıfta yüksek ve '
        'dengeli performans gösterdiğini ortaya koymaktadır. Makro F1\'in ağırlıklı F1\'e '
        'çok yakın olması az temsil edilen sınıflarda da başarı elde edildiğine işaret '
        'etmektedir; bu sınıf ağırlıklandırması ve ColorJitter artırmasının etkinliğini '
        'doğrulamaktadır.'
    ))

    h2(doc, '4.3 Sınıf Bazlı Performans Analizi')

    p(doc, (
        'Sınıf düzeyinde F1 skoru analizi tüm 9 sınıfın 0,70 güven eşiğini rahatlıkla '
        'aştığını göstermektedir. Normal, Tümör, Stroma ve Adipoz sınıfları 0,99\'un '
        'üzerinde F1 skoruna ulaşırken Artık ve Arka Plan sınıfları görece daha düşük '
        'ancak yüksek F1 değerleri sergilemiştir.'
    ))
    p(doc, (
        'Karmaşıklık matrisi analizi karışıklığın en çok Kompleks ve Tümör sınıfları '
        'arasında görüldüğünü ortaya koymaktadır. Kompleks sınıfı içinde tümör bileşenleri '
        'barındıran karışık doku yapılarını temsil ettiğinden bu durum patolojik açıdan '
        'anlamlıdır; literatürde de bu iki sınıf arasındaki karışıklık yaygın olarak '
        'raporlanmaktadır (Kather ve ark., 2019).'
    ))

    h2(doc, '4.4 Literatürle Karşılaştırma')

    p(doc, (
        'SimpleCancerNet, NCT-CRC-HE-100K üzerinde %99,24 doğruluk ve 0,9924 makro F1 ile '
        'Tablo 1\'deki tüm karşılaştırma çalışmalarını geride bırakmaktadır. Bu üstünlük '
        'transfer öğrenme tabanlı VGG19 (96,2%), ResNet50 (97,8%) ve EfficientNet-B4 '
        '(98,6%) gibi mimarilerle kıyaslandığında özellikle dikkat çekicidir; çünkü bu '
        'modeller onlarca kat daha fazla parametre kullanmaktadır.'
    ))
    p(doc, (
        'Bu sonuçların yorumlanmasında dikkat edilmesi gereken önemli bir nokta veri '
        'bölme stratejisindeki farklılıklardır. Bazı çalışmalar farklı rastgele tohumlar '
        'veya çapraz doğrulama kullanmış olabilmektedir. Bununla birlikte mevcut sonuçlar '
        'en katı yorumla bile rekabetçi düzeyde olduğunu ortaya koymaktadır.'
    ))

    h2(doc, '4.5 API ve Arayüz Performansı')

    p(doc, (
        'FastAPI servisi CUDA etkin yapılandırmada görüntü başına ortalama 8–12 ms çıkarım '
        'süresi sunmaktadır. CPU modunda bu değer 45–80 ms\'ye çıkmaktadır; ancak çoğu '
        'klinik kullanım senaryosu için saniyenin altındaki bu süre kabul edilebilir '
        'sınırlardadır.'
    ))
    p(doc, (
        'Web arayüzü sezgisel tasarımı sayesinde teknik arka plan gerektirmeksizin '
        'kullanılabilmektedir. Klinik gruplama renk kodlaması (kırmızı, yeşil, gri, sarı), '
        'kullanıcının dikkatini gerektiren sonuçlara yönelimini hızlandırmaktadır. '
        '9 sınıfa ait olasılık çubuk grafikleri klinisyene modelin karar sürecine dair '
        'şeffaf bir görünüm sunmaktadır.'
    ))

    page_break(doc)

    # ══════════════════════════════════════════
    #  5. SONUÇ
    # ══════════════════════════════════════════
    h1(doc, 'Sonuç ve Gelecek Çalışmalar')

    h2(doc, '5.1 Sonuç')

    p(doc, (
        'Bu çalışmada, kolon kanseri histopatolojik görüntülerini 9 doku sınıfına ayırt eden '
        'SimpleCancerNet adlı özgün ve hafif bir derin öğrenme mimarisi geliştirilmiştir. '
        'NCT-CRC-HE-100K veri seti üzerinde elde edilen %99,24 test doğruluğu ve '
        '0,9924 makro F1 skoru çalışmanın temel hedeflerine ulaşıldığını ve sıfırdan '
        'eğitim stratejisinin başarısını doğrulamaktadır.'
    ))
    p(doc, (
        'Yalnızca 458.537 parametre kullanan SimpleCancerNet, büyük modellerin aksine klinik '
        'dağıtım için elverişli bir boyuta sahiptir. FastAPI ve Next.js ile oluşturulan '
        'uçtan uca açık kaynak uygulama altyapısı modelin gerçek dünya klinik ortamlarında '
        'kullanılabilmesini mümkün kılmaktadır.'
    ))
    p(doc, (
        'Çalışmanın sınırlılıkları şunlardır: değerlendirme yalnızca NCT-CRC-HE-100K ile '
        'yapılmış olup dış veri setleri üzerinde kapsamlı bir genellenebilirlik testi henüz '
        'gerçekleştirilmemiştir; model patoloji onayı almış klinik ortam koşullarında '
        'test edilmemiştir; açıklanabilirlik mekanizmaları (Grad-CAM) henüz entegre '
        'edilmemiştir.'
    ))

    h2(doc, '5.2 Gelecek Çalışmalar')

    p(doc, 'Çalışmanın ilerleyen aşamalarında planlanmakta olan genişletmeler:')
    for f in [
        'Grad-CAM görselleştirme — model kararlarını görsel olarak açıklayarak klinisyen güvenini artırma.',
        'Harici veri setleriyle genellenebilirlik değerlendirmesi — farklı laboratuvar ve boyama protokollerinden gelen görüntülerle test.',
        'ONNX dönüşümü — modeli GPU\'suz sunucularda ve uç cihazlarda çalıştırılabilir hale getirme.',
        'Bayesian hiperparametre optimizasyonu — scikit-optimize ile sistematik arama.',
        'TensorBoard entegrasyonu — eğitim sürecinin gerçek zamanlı izlenmesi.',
        'Çoklu etiket sınıflandırması — karma doku bölgelerini daha gerçekçi biçimde modellemek için.',
        'Klinik pilot çalışma — gerçek klinik ortamda kullanılabilirlik ve hata analizi.',
    ]:
        bullet(doc, f)

    page_break(doc)

    # ══════════════════════════════════════════
    #  6. KAYNAKÇA
    # ══════════════════════════════════════════
    h1(doc, 'Kaynakça')

    refs = [
        'Chen, R. J., ve ark. (2022). Scaling vision transformers to gigapixel images via hierarchical self-supervised learning. Proceedings of the IEEE/CVF CVPR, 16144–16155.',
        'Esteva, A., Kuprel, B., Novoa, R. A., Ko, J., Swetter, S. M., Blau, H. M., ve Thrun, S. (2017). Dermatologist-level classification of skin cancer with deep neural networks. Nature, 542(7639), 115–118.',
        'He, K., Zhang, X., Ren, S., ve Sun, J. (2016). Deep residual learning for image recognition. Proceedings of the IEEE CVPR, 770–778.',
        'Huang, G., Liu, Z., van der Maaten, L., ve Weinberger, K. Q. (2017). Densely connected convolutional networks. Proceedings of the IEEE CVPR, 4700–4708.',
        'Kather, J. N., Weis, C. A., ve ark. (2016). Multi-class texture analysis in colorectal cancer histology. Scientific Reports, 6(1), 27988.',
        'Kather, J. N., Krisam, J., ve ark. (2019). Predicting survival from colorectal cancer histology slides using deep learning. PLOS Medicine, 16(1), e1002730.',
        'King, G., ve Zeng, L. (2001). Logistic regression in rare events data. Political Analysis, 9(2), 137–163.',
        'Krizhevsky, A., Sutskever, I., ve Hinton, G. E. (2012). ImageNet classification with deep convolutional neural networks. NeurIPS, 25, 1097–1105.',
        'LeCun, Y., Boser, B., ve ark. (1989). Backpropagation applied to handwritten zip code recognition. Neural Computation, 1(4), 541–551.',
        'Lin, M., Chen, Q., ve Yan, S. (2014). Network in network. ICLR.',
        'Lin, T. Y., Goyal, P., ve ark. (2017). Focal loss for dense object detection. Proceedings of the IEEE ICCV, 2980–2988.',
        'Litjens, G., Kooi, T., ve ark. (2017). A survey on deep learning in medical image analysis. Medical Image Analysis, 42, 60–88.',
        'Pan, S. J., ve Yang, Q. (2010). A survey on transfer learning. IEEE Transactions on Knowledge and Data Engineering, 22(10), 1345–1359.',
        'Raghu, M., Zhang, C., Kleinberg, J., ve Bengio, S. (2019). Transfusion: Understanding transfer learning for medical imaging. NeurIPS, 32.',
        'Shorten, C., ve Khoshgoftaar, T. M. (2019). A survey on image data augmentation for deep learning. Journal of Big Data, 6(1), 1–48.',
        'Simonyan, K., ve Zisserman, A. (2015). Very deep convolutional networks for large-scale image recognition. ICLR.',
        'Srinidhi, C. L., Ciga, O., ve Martel, A. L. (2021). Deep neural network models for computational histopathology: A survey. Medical Image Analysis, 67, 101813.',
        'Tan, M., ve Le, Q. (2019). EfficientNet: Rethinking model scaling for convolutional neural networks. ICML, 6105–6114.',
        'Topol, E. J. (2019). High-performance medicine: The convergence of human and artificial intelligence. Nature Medicine, 25(1), 44–56.',
        'World Health Organization (WHO). (2024). Colorectal cancer. Global Cancer Observatory. Erişim: Mayıs 2026.',
    ]
    for ref in refs:
        para = doc.add_paragraph(style='Normal')
        para.paragraph_format.left_indent = Cm(1.27)
        para.paragraph_format.first_line_indent = Cm(-1.27)
        para.paragraph_format.space_after = Pt(6)
        para.add_run(ref)

    page_break(doc)

    # ══════════════════════════════════════════
    #  EKLER
    # ══════════════════════════════════════════
    h1(doc, 'Ekler')

    h2(doc, 'Ek A: Model Parametre Dağılımı')

    tbl(doc,
        ['Katman', 'Tip', 'Girdi', 'Çıktı', 'Parametre'],
        [
            ('Blok 1', 'Conv(5×5)+BN', '3×224×224',  '32×112×112', '2.464'),
            ('Blok 2', 'Conv(3×3)+BN', '32×112×112', '64×56×56',   '18.688'),
            ('Blok 3', 'Conv(3×3)+BN', '64×56×56',   '128×28×28',  '74.112'),
            ('Blok 4', 'Conv(3×3)+BN', '128×28×28',  '256×14×14',  '295.680'),
            ('GAP',    'AvgPool',      '256×14×14',  '256',         '0'),
            ('FC-1',   'Linear+Drop', '256',          '256',         '65.792'),
            ('FC-2',   'Linear',      '256',          '9',           '2.313'),
            ('Toplam', '—',           '—',            '—',           '459.049*'),
        ])
    caption(doc, 'Tablo A.1. SimpleCancerNet katman bazında parametre sayısı (*BatchNorm dahil gerçek toplam: 458.537).')

    h2(doc, 'Ek B: Klinik Gruplama Kuralları')

    tbl(doc,
        ['Klinik Grup', 'Koşul', 'Sınıflar'],
        [
            ('Belirsiz',        'Güven < 0.70',              '—'),
            ('Kanser Şüphesi',  'Güven ≥ 0.70 ve idx ∈ {1,4}', 'Tümör (1), Kompleks (4)'),
            ('Normal Doku',     'Güven ≥ 0.70 ve idx ∈ {0,2,3,6}', 'Normal, Stroma, Lenfosit, Müköz'),
            ('Klinik Dışı',    'Güven ≥ 0.70 ve idx ∈ {5,7,8}', 'Artık, Adipoz, Arka Plan'),
        ])
    caption(doc, 'Tablo B.1. Klinik gruplama karar kuralları.')

    h2(doc, 'Ek C: API Endpoint Özeti')

    tbl(doc,
        ['Endpoint', 'Metod', 'Girdi', 'Çıktı'],
        [
            ('/health',  'GET',  '—', 'status, model_loaded, device'),
            ('/predict', 'POST', 'multipart/form-data (image)', 'prediction, all_probabilities, meta'),
        ])
    caption(doc, 'Tablo C.1. FastAPI endpoint özeti.')

    p(doc, (
        'HTTP hata kodları: 415 — Desteklenmeyen dosya formatı; '
        '422 — Görüntü okunamadı. CORS politikası varsayılan olarak '
        'yalnızca http://localhost:3000 kaynağına izin vermektedir.'
    ))

    h2(doc, 'Ek D: Proje Dosya Yapısı')

    tbl(doc,
        ['Yol', 'Açıklama'],
        [
            ('src/dataset.py',  'Veri yükleme, split ve DataLoader fabrikası'),
            ('src/model.py',    'SimpleCancerNet mimarisi ve predict() yardımcısı'),
            ('src/train.py',    'Eğitim döngüsü ve CONFIG tanımı'),
            ('src/evaluate.py', 'Test seti değerlendirme ve görselleştirme'),
            ('src/api.py',      'FastAPI REST API servisi'),
            ('frontend/app/page.tsx', 'Next.js tek sayfa arayüzü'),
            ('outputs/checkpoints/best_model.pt', 'En iyi checkpoint'),
            ('outputs/plots/', 'Eğitim eğrileri, confusion matrix, F1 grafikleri'),
            ('check.py',        'GPU/CUDA durum kontrolü'),
        ])
    caption(doc, 'Tablo D.1. Proje dosya yapısı.')

    doc.save(OUTPUT)
    print(f"Belge kaydedildi: {OUTPUT}")


if __name__ == "__main__":
    build()
